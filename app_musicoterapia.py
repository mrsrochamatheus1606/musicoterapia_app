import os
import math
from io import BytesIO
from datetime import date

import streamlit as st
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt


# ==========================================================
# CONFIGURAÇÃO GERAL DO APP
# ==========================================================
st.set_page_config(
    page_title="Pulsar MT System",
    page_icon="logo.png" if os.path.exists("logo.png") else "🎵",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ==========================================================
# CSS PREMIUM
# ==========================================================
st.markdown("""
<style>

/* Fundo geral */
[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #020617 0%, #0f172a 45%, #111827 100%);
    color: #f8fafc;
}

/* Container principal */
.block-container {
    padding-top: 1.8rem;
    padding-bottom: 3rem;
    max-width: 1400px;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #020617 0%, #0f172a 100%);
    border-right: 1px solid rgba(148, 163, 184, 0.18);
}

[data-testid="stSidebar"] * {
    color: #e2e8f0;
}

/* Títulos */
h1, h2, h3 {
    color: #e0f2fe;
    font-family: 'Segoe UI', sans-serif;
}

h1 {
    font-size: 2.5rem !important;
    font-weight: 800 !important;
}

h2 {
    font-size: 1.55rem !important;
    font-weight: 750 !important;
    margin-top: 1.5rem !important;
}

h3 {
    font-size: 1.15rem !important;
    font-weight: 700 !important;
}

/* Texto */
p, label, span, div {
    font-family: 'Segoe UI', sans-serif;
}

/* Inputs */
.stTextInput input,
.stNumberInput input,
.stTextArea textarea,
.stSelectbox div[data-baseweb="select"],
.stMultiSelect div[data-baseweb="select"] {
    border-radius: 14px !important;
    border: 1px solid rgba(148, 163, 184, 0.28) !important;
    background-color: rgba(15, 23, 42, 0.85) !important;
    color: #f8fafc !important;
}

.stTextInput input:focus,
.stNumberInput input:focus,
.stTextArea textarea:focus {
    border-color: #38bdf8 !important;
    box-shadow: 0 0 0 1px #38bdf8 !important;
}

/* Radio */
.stRadio > div {
    background: rgba(15, 23, 42, 0.55);
    border-radius: 14px;
    padding: 0.55rem 0.8rem;
    border: 1px solid rgba(148, 163, 184, 0.15);
}

/* Botão principal */
.stButton > button {
    width: 100%;
    min-height: 3.2rem;
    border-radius: 16px;
    border: none;
    background: linear-gradient(90deg, #06b6d4, #6366f1, #8b5cf6);
    color: white;
    font-weight: 800;
    font-size: 1rem;
    box-shadow: 0 14px 35px rgba(37, 99, 235, 0.25);
    transition: all 0.2s ease-in-out;
}

.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 18px 42px rgba(37, 99, 235, 0.38);
    color: white;
}

/* Download button */
.stDownloadButton > button {
    width: 100%;
    min-height: 3.2rem;
    border-radius: 16px;
    border: none;
    background: linear-gradient(90deg, #10b981, #06b6d4);
    color: white;
    font-weight: 800;
    font-size: 1rem;
    box-shadow: 0 14px 35px rgba(16, 185, 129, 0.25);
}

/* Abas */
.stTabs [data-baseweb="tab-list"] {
    gap: 10px;
}

.stTabs [data-baseweb="tab"] {
    background-color: rgba(15, 23, 42, 0.95);
    border-radius: 14px 14px 0 0;
    padding: 12px 18px;
    border: 1px solid rgba(148, 163, 184, 0.15);
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(90deg, rgba(14, 165, 233, 0.35), rgba(99, 102, 241, 0.35));
    border-bottom: 2px solid #38bdf8;
}

/* Métricas */
[data-testid="stMetric"] {
    background: rgba(15, 23, 42, 0.72);
    border: 1px solid rgba(148, 163, 184, 0.18);
    padding: 1rem;
    border-radius: 18px;
    box-shadow: 0 14px 35px rgba(2, 6, 23, 0.28);
}

/* Expander */
.streamlit-expanderHeader {
    background: rgba(15, 23, 42, 0.75) !important;
    border-radius: 14px !important;
    color: #e0f2fe !important;
}

/* Alertas */
[data-testid="stAlert"] {
    border-radius: 16px;
    border: 1px solid rgba(148, 163, 184, 0.18);
}

/* Linha divisória */
hr {
    border: none;
    height: 1px;
    background: rgba(148, 163, 184, 0.2);
    margin: 1.5rem 0;
}

/* Cards customizados */
.premium-card {
    background: rgba(15, 23, 42, 0.76);
    border: 1px solid rgba(148, 163, 184, 0.18);
    border-radius: 22px;
    padding: 22px;
    box-shadow: 0 18px 45px rgba(2, 6, 23, 0.30);
    margin-bottom: 18px;
}

.hero-card {
    background: radial-gradient(circle at top left, rgba(6, 182, 212, 0.22), transparent 35%),
                linear-gradient(135deg, rgba(15, 23, 42, 0.94), rgba(30, 41, 59, 0.82));
    border: 1px solid rgba(56, 189, 248, 0.25);
    border-radius: 28px;
    padding: 30px;
    box-shadow: 0 22px 60px rgba(2, 6, 23, 0.45);
    margin-bottom: 22px;
}

.hero-title {
    font-size: 2.55rem;
    font-weight: 900;
    letter-spacing: -0.04em;
    color: #f8fafc;
    margin-bottom: 0.35rem;
}

.hero-subtitle {
    font-size: 1.05rem;
    color: #cbd5e1;
    margin-bottom: 0.75rem;
}

.pill {
    display: inline-block;
    padding: 7px 12px;
    border-radius: 999px;
    background: rgba(14, 165, 233, 0.14);
    color: #bae6fd;
    border: 1px solid rgba(56, 189, 248, 0.25);
    margin-right: 8px;
    font-size: 0.85rem;
    font-weight: 700;
}

.section-label {
    font-size: 0.82rem;
    color: #67e8f9;
    text-transform: uppercase;
    letter-spacing: 0.12em;
    font-weight: 800;
    margin-bottom: 0.3rem;
}

.small-muted {
    color: #94a3b8;
    font-size: 0.9rem;
}

</style>
""", unsafe_allow_html=True)


# ==========================================================
# SIDEBAR / MARCA
# ==========================================================
with st.sidebar:
    if os.path.exists("logo.png"):
        st.markdown("<div style='text-align:center; padding-top:10px;'>", unsafe_allow_html=True)
        st.image("logo.png", width=165)
        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="text-align:center; padding: 10px 4px 18px 4px;">
            <div style="font-size:2.2rem;">🎵</div>
            <h2 style="margin-bottom:0;">Pulsar MT</h2>
            <p style="color:#94a3b8; margin-top:0;">Clinical System</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("**Sistema para:**")
    st.markdown("- Análise de dados musicoterapêuticos")
    st.markdown("- Plano terapêutico automatizado")
    st.markdown("- GAS / SMART automático")
    st.markdown("- Relatório clínico em Word")
    st.markdown("---")
    st.info("Preencha as etapas e gere um relatório técnico em poucos minutos.")


# ==========================================================
# HEADER PREMIUM
# ==========================================================
if os.path.exists("logo.png"):
    st.markdown("""
    <div class="hero-card">
        <div style="display:flex; align-items:center; gap:22px; flex-wrap:wrap;">
            <div style="background:rgba(255,255,255,0.06); border:1px solid rgba(148,163,184,0.18); border-radius:22px; padding:12px;">
    """, unsafe_allow_html=True)
    st.image("logo.png", width=92)
    st.markdown("""
            </div>
            <div>
                <div class="section-label">Musicoterapia baseada em dados</div>
                <div class="hero-title">Pulsar MT System</div>
                <div class="hero-subtitle">
                    Análise inteligente de avaliações, geração automática de objetivos clínicos, plano terapêutico, GAS e relatório profissional.
                </div>
                <span class="pill">📊 Dados clínicos</span>
                <span class="pill">🎼 Musicoterapia</span>
                <span class="pill">📄 Relatórios rápidos</span>
                <span class="pill">🧠 Plano terapêutico</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div class="hero-card">
        <div class="section-label">Musicoterapia baseada em dados</div>
        <div class="hero-title">Pulsar MT System</div>
        <div class="hero-subtitle">
            Análise inteligente de avaliações, geração automática de objetivos clínicos, plano terapêutico, GAS e relatório profissional.
        </div>
        <span class="pill">📊 Dados clínicos</span>
        <span class="pill">🎼 Musicoterapia</span>
        <span class="pill">📄 Relatórios rápidos</span>
        <span class="pill">🧠 Plano terapêutico</span>
    </div>
    """, unsafe_allow_html=True)


# ==========================================================
# FUNÇÕES VISUAIS
# ==========================================================
def card_inicio(titulo, subtitulo=None):
    st.markdown("<div class='premium-card'>", unsafe_allow_html=True)
    st.markdown(f"### {titulo}")
    if subtitulo:
        st.markdown(f"<p class='small-muted'>{subtitulo}</p>", unsafe_allow_html=True)


def card_fim():
    st.markdown("</div>", unsafe_allow_html=True)


# ==========================================================
# SESSION STATE
# ==========================================================
if "relatorio_gerado" not in st.session_state:
    st.session_state.relatorio_gerado = False


# ==========================================================
# ABAS PRINCIPAIS
# ==========================================================
aba_paciente, aba_avaliacoes, aba_resultados, aba_relatorio = st.tabs([
    "👤 Paciente",
    "🎼 Avaliações",
    "📊 Análise automática",
    "📄 Relatório"
])


# ==========================================================
# ABA 1 — PACIENTE
# ==========================================================
with aba_paciente:
    card_inicio("👤 Dados do Profissional", "Identificação do responsável técnico pelo relatório.")
    col1, col2 = st.columns(2)
    with col1:
        terapeuta = st.text_input("Nome completo do terapeuta")
    with col2:
        registro = st.text_input("Registro profissional")
    card_fim()

    card_inicio("🧒 Identificação do Paciente", "Dados clínicos e contextuais do paciente avaliado.")
    col1, col2, col3 = st.columns(3)
    with col1:
        nome = st.text_input("Nome do paciente")
    with col2:
        idade = st.number_input("Idade", min_value=0, max_value=120, value=0, step=1)
    with col3:
        diagnostico = st.text_input("Diagnóstico")

    col4, col5, col6 = st.columns(3)
    with col4:
        data_nascimento = st.text_input("Data de nascimento")
    with col5:
        escolaridade = st.text_input("Escolaridade")
    with col6:
        responsaveis = st.text_input("Pais ou responsáveis")

    col7, col8, col9 = st.columns(3)
    with col7:
        metodo_intensivo = st.selectbox(
            "Método / abordagem",
            ["", "MIG", "TREINI", "ABA", "Particular"]
        )
    with col8:
        data_inicio = st.text_input("Data de início da intervenção")
    with col9:
        data_atual = st.text_input("Data atual", value=date.today().strftime("%d/%m/%Y"))
    card_fim()

    card_inicio("🧾 Informações Clínicas", "Campos narrativos que irão enriquecer o parecer técnico.")
    historia_clinica = st.text_area("História clínica", height=130)
    queixa = st.text_area("Queixa principal / motivo do encaminhamento", height=100)
    observacoes = st.text_area("Observações clínicas", height=100)
    preferencias = st.text_area("Preferências e recusas sonoro-musicais", height=100)
    interacoes = st.text_area("Interações sonoro-musicais", height=100)
    card_fim()


# ==========================================================
# VARIÁVEIS DAS ESCALAS
# ==========================================================
def campo(label, key):
    return st.number_input(label, min_value=0, max_value=5, value=0, step=1, key=key)


nordoff_dominios = {}
iaps_improvisacao = {}
iaps_recriacao = {}
iaps_composicao = {}
iaps_escuta = {}
demuca_respostas = {}


DEMUCA_DOMINIOS = {
    "Comportamentos Restritivos": {
        "tipo": "restritivo",
        "maximo": 14,
        "itens": [
            ("Estereotipias", 1),
            ("Agressividade", 1),
            ("Desinteresse", 1),
            ("Passividade", 1),
            ("Reclusão (Isolamento)", 1),
            ("Resistência", 1),
            ("Pirraça", 1),
        ],
    },
    "Interação Social - Cognição": {
        "tipo": "positivo",
        "maximo": 18,
        "itens": [
            ("Contato visual", 1),
            ("Comunicação verbal", 1),
            ("Interação com outros objetos", 1),
            ("Interação com instrumentos musicais", 1),
            ("Interação com educador/musicoterapeuta", 1),
            ("Interação com os pais (se aplicável)", 1),
            ("Interação com os pares (se aplicável)", 1),
            ("Atenção", 1),
            ("Imitação", 1),
        ],
    },
    "Percepção - Exploração Rítmica": {
        "tipo": "positivo",
        "maximo": 16,
        "itens": [
            ("Pulso Interno", 1),
            ("Regulação Temporal", 1),
            ("Ritmo Real", 2),
            ("Apoio", 2),
            ("Contrastes de andamento", 2),
        ],
    },
    "Percepção - Exploração Sonora": {
        "tipo": "positivo",
        "maximo": 14,
        "itens": [
            ("Som / Silêncio", 1),
            ("Timbre", 1),
            ("Planos de altura", 1),
            ("Movimento sonoro", 1),
            ("Contrastes de intensidade", 1),
            ("Repetição de ideias rítmicas e/ou melódicas", 1),
            ("Senso de conclusão", 1),
        ],
    },
    "Exploração Vocal": {
        "tipo": "positivo",
        "maximo": 14,
        "itens": [
            ("Vocalizações", 1),
            ("Balbucios", 1),
            ("Sílabas canônicas", 1),
            ("Imitação de canções", 2),
            ("Criação vocal", 2),
        ],
    },
    "Movimentação corporal com a música": {
        "tipo": "positivo",
        "maximo": 14,
        "itens": [
            ("Andar", 1),
            ("Correr", 1),
            ("Parar", 1),
            ("Gesticular", 1),
            ("Dançar", 1),
            ("Movimentar-se no lugar", 1),
            ("Pular", 1),
        ],
    },
}


# ==========================================================
# ABA 2 — AVALIAÇÕES
# ==========================================================
with aba_avaliacoes:
    card_inicio("🎼 Seleção de Escalas", "Escolha apenas os instrumentos que serão aplicados. O relatório explicará somente as escalas selecionadas.")
    escalas_escolhidas = st.multiselect(
        "Selecione as avaliações",
        ["Nordoff-Robbins", "IAPS", "DEMUCA"],
        default=["Nordoff-Robbins", "IAPS"]
    )
    st.success("A GAS, o plano terapêutico, a prescrição e a conduta serão gerados automaticamente a partir dos scores.")
    card_fim()

    if "Nordoff-Robbins" in escalas_escolhidas:
        card_inicio("🎵 Escala Nordoff-Robbins", "Pontuação de 0 a 5 para cada domínio de comunicabilidade musical.")
        col1, col2 = st.columns(2)
        with col1:
            nordoff_dominios["Expressão emocional"] = campo("Expressão emocional", "n_exp_emocional")
            nordoff_dominios["Exploração sonora"] = campo("Exploração sonora", "n_exploracao_sonora")
            nordoff_dominios["Interação musical"] = campo("Interação musical", "n_interacao_musical")
            nordoff_dominios["Engajamento"] = campo("Engajamento", "n_engajamento")
            nordoff_dominios["Responsividade musical"] = campo("Responsividade musical", "n_responsividade")
        with col2:
            nordoff_dominios["Iniciativa musical"] = campo("Iniciativa musical", "n_iniciativa")
            nordoff_dominios["Sustentação da atividade musical"] = campo("Sustentação da atividade musical", "n_sustentacao")
            nordoff_dominios["Comunicação não verbal"] = campo("Comunicação não verbal", "n_comunicacao_nao_verbal")
            nordoff_dominios["Reciprocidade musical"] = campo("Reciprocidade musical", "n_reciprocidade")
            nordoff_dominios["Organização musical"] = campo("Organização musical", "n_organizacao")
        card_fim()

    if "IAPS" in escalas_escolhidas:
        card_inicio("🎹 IAPS", "Avaliação por áreas: improvisação, recriação, composição e escuta musical.")
        col1, col2 = st.columns(2)

        with col1:
            with st.expander("IAPS - Improvisação", expanded=True):
                iaps_improvisacao = {
                    "Iniciativa sonora": campo("Iniciativa sonora", "i_iniciativa"),
                    "Resposta musical": campo("Resposta musical", "i_resposta"),
                    "Organização sonora": campo("Organização sonora", "i_organizacao"),
                    "Interação musical": campo("Interação musical", "i_interacao"),
                }

            with st.expander("IAPS - Recriação", expanded=True):
                iaps_recriacao = {
                    "Memória musical": campo("Memória musical", "r_memoria"),
                    "Coordenação motora": campo("Coordenação motora", "r_coordenacao"),
                    "Seguimento musical": campo("Seguimento musical", "r_seguimento"),
                    "Participação": campo("Participação", "r_participacao"),
                }

        with col2:
            with st.expander("IAPS - Composição", expanded=True):
                iaps_composicao = {
                    "Criatividade": campo("Criatividade", "c_criatividade"),
                    "Organização de ideias": campo("Organização de ideias", "c_organizacao"),
                    "Expressão simbólica": campo("Expressão simbólica", "c_expressao"),
                    "Autoria": campo("Autoria", "c_autoria"),
                }

            with st.expander("IAPS - Escuta Musical", expanded=True):
                iaps_escuta = {
                    "Atenção auditiva": campo("Atenção auditiva", "e_atencao"),
                    "Resposta emocional": campo("Resposta emocional", "e_resposta"),
                    "Reflexão": campo("Reflexão", "e_reflexao"),
                    "Integração da experiência sonora": campo("Integração da experiência sonora", "e_integracao"),
                }
        card_fim()

    if "DEMUCA" in escalas_escolhidas:
        card_inicio("🥁 DEMUCA", "Escala de classificação: N = Não | P = Pouco | M = Muito.")
        for dominio, info in DEMUCA_DOMINIOS.items():
            with st.expander(f"DEMUCA - {dominio}", expanded=False):
                demuca_respostas[dominio] = {}
                for item, peso in info["itens"]:
                    label = f"{item}" + ("  (x2)" if peso == 2 else "")
                    resposta = st.radio(
                        label,
                        ["N", "P", "M"],
                        horizontal=True,
                        key=f"demuca_{dominio}_{item}"
                    )
                    demuca_respostas[dominio][item] = {"resposta": resposta, "peso": peso}
        card_fim()


# ==========================================================
# FUNÇÕES DE CÁLCULO
# ==========================================================
def classificar(valor, maximo):
    percentual_calc = (valor / maximo) * 100 if maximo else 0
    if percentual_calc < 40:
        return "baixo"
    elif percentual_calc < 70:
        return "moderado"
    return "adequado"


def percentual(valor, maximo):
    return round((valor / maximo) * 100, 1) if maximo else 0


def totals_to_order(totais):
    ordem = ["Improvisação", "Recriação", "Composição", "Escuta Musical"]
    return {k: totais[k] for k in ordem if k in totais}


def calcular_demuca(demuca_respostas):
    totais = {}
    for dominio, itens in demuca_respostas.items():
        tipo = DEMUCA_DOMINIOS[dominio]["tipo"]
        total = 0
        for item, dados in itens.items():
            resposta = dados["resposta"]
            peso = dados["peso"]
            mapa = {"N": 2, "P": 1, "M": 0} if tipo == "restritivo" else {"N": 0, "P": 1, "M": 2}
            total += mapa[resposta] * peso
        totais[dominio] = total
    return totais


def identificar_prejuizos(nordoff_total, totais_iaps, totais_demuca):
    prejuizos = []

    if nordoff_total is not None:
        nivel = classificar(nordoff_total, 50)
        if nivel in ["baixo", "moderado"]:
            prejuizos.append({
                "area": "Comunicabilidade musical",
                "origem": "Nordoff-Robbins",
                "nivel": nivel,
                "objetivo": "Ampliar a comunicação musical funcional, favorecendo iniciativa, responsividade, reciprocidade e sustentação da interação sonoro-musical.",
                "habilidade": "comunicação musical, responsividade e vínculo terapêutico"
            })

    if totais_iaps:
        mapa_iaps = {
            "Improvisação": {
                "objetivo": "Estimular a iniciativa sonora, a espontaneidade musical e a construção de respostas musicais intencionais em contexto improvisacional.",
                "habilidade": "iniciativa sonora e expressão espontânea"
            },
            "Recriação": {
                "objetivo": "Fortalecer memória musical, coordenação motora, imitação, seguimento de modelos e participação em atividades musicais estruturadas.",
                "habilidade": "memória musical, coordenação e participação estruturada"
            },
            "Composição": {
                "objetivo": "Favorecer criatividade, autoria, simbolização e organização de ideias musicais por meio da produção musical guiada.",
                "habilidade": "criatividade, autoria e elaboração simbólica"
            },
            "Escuta Musical": {
                "objetivo": "Promover atenção auditiva, escuta ativa, resposta emocional e integração subjetiva da experiência sonora.",
                "habilidade": "escuta ativa, atenção e integração sonora"
            },
        }
        for area, valor in totals_to_order(totais_iaps).items():
            nivel = classificar(valor, 20)
            if nivel in ["baixo", "moderado"]:
                prejuizos.append({
                    "area": area,
                    "origem": "IAPS",
                    "nivel": nivel,
                    "objetivo": mapa_iaps[area]["objetivo"],
                    "habilidade": mapa_iaps[area]["habilidade"]
                })

    if totais_demuca:
        mapa_demuca = {
            "Comportamentos Restritivos": {
                "objetivo": "Reduzir a interferência de comportamentos restritivos no setting musicoterapêutico, favorecendo co-regulação, previsibilidade, disponibilidade relacional e engajamento musical funcional.",
                "habilidade": "regulação, disponibilidade relacional e redução de interferências comportamentais"
            },
            "Interação Social - Cognição": {
                "objetivo": "Ampliar contato visual, atenção compartilhada, imitação, comunicação e interação social mediada pela música.",
                "habilidade": "interação social, atenção compartilhada e cognição musical"
            },
            "Percepção - Exploração Rítmica": {
                "objetivo": "Desenvolver pulso interno, regulação temporal, apoio rítmico, ritmo real e percepção de contrastes de andamento.",
                "habilidade": "percepção rítmica, organização temporal e coordenação musical"
            },
            "Percepção - Exploração Sonora": {
                "objetivo": "Estimular discriminação sonora, percepção de timbre, planos de altura, movimento sonoro, intensidade, repetição de ideias musicais e senso de conclusão.",
                "habilidade": "percepção sonora, discriminação auditiva e organização musical"
            },
            "Exploração Vocal": {
                "objetivo": "Favorecer vocalizações, balbucios, sílabas canônicas, imitação de canções e criação vocal em contexto terapêutico.",
                "habilidade": "expressão vocal, comunicação vocal e musicalidade da voz"
            },
            "Movimentação corporal com a música": {
                "objetivo": "Ampliar organização motora, expressão corporal, coordenação, deslocamento e movimentação funcional associada à experiência musical.",
                "habilidade": "expressão corporal, coordenação motora e integração música-movimento"
            },
        }
        for dominio, valor in totais_demuca.items():
            nivel = classificar(valor, DEMUCA_DOMINIOS[dominio]["maximo"])
            if nivel in ["baixo", "moderado"]:
                prejuizos.append({
                    "area": dominio,
                    "origem": "DEMUCA",
                    "nivel": nivel,
                    "objetivo": mapa_demuca[dominio]["objetivo"],
                    "habilidade": mapa_demuca[dominio]["habilidade"]
                })

    if not prejuizos:
        prejuizos.append({
            "area": "Manutenção e ampliação de repertório musicoterapêutico",
            "origem": "Síntese clínica",
            "nivel": "adequado",
            "objetivo": "Aprofundar os recursos musicais já estabelecidos, ampliando complexidade expressiva, autonomia, flexibilidade interacional e elaboração simbólica no setting musicoterapêutico.",
            "habilidade": "autonomia musical, flexibilidade e elaboração expressiva"
        })

    return prejuizos


# ==========================================================
# GRÁFICOS
# ==========================================================
def gerar_grafico_barras(titulo, dados, maximos=None):
    categorias = list(dados.keys())
    valores = list(dados.values())
    limite = max(maximos.values()) if maximos else (max(valores) if valores else 5)

    fig, ax = plt.subplots(figsize=(11, 5.5))
    fig.patch.set_facecolor("#020617")
    ax.set_facecolor("#0f172a")
    ax.bar(categorias, valores)
    ax.set_title(titulo, color="white", fontsize=14, fontweight="bold")
    ax.set_ylabel("Pontuação", color="white")
    ax.set_ylim(0, limite)
    ax.tick_params(axis="x", rotation=35, colors="white")
    ax.tick_params(axis="y", colors="white")
    ax.spines["bottom"].set_color("#475569")
    ax.spines["left"].set_color("#475569")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()

    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight", dpi=180, facecolor=fig.get_facecolor())
    plt.close(fig)
    buffer.seek(0)
    return buffer


def gerar_grafico_iaps(totais_iaps):
    return gerar_grafico_barras(
        "IAPS - Pontuação por área",
        totais_iaps,
        {"Improvisação": 20, "Recriação": 20, "Composição": 20, "Escuta Musical": 20}
    )


def gerar_grafico_demuca(totais_demuca):
    maximos = {dominio: DEMUCA_DOMINIOS[dominio]["maximo"] for dominio in totais_demuca}
    return gerar_grafico_barras("DEMUCA - Pontuação por domínio", totais_demuca, maximos)


def gerar_grafico_radar(totais_iaps, nordoff_total=None, totais_demuca=None):
    categorias = []
    valores_percentuais = []

    if nordoff_total is not None:
        categorias.append("Nordoff")
        valores_percentuais.append(percentual(nordoff_total, 50))

    if totais_iaps:
        categorias.extend(["Improvisação", "Recriação", "Composição", "Escuta"])
        valores_percentuais.extend([
            percentual(totais_iaps["Improvisação"], 20),
            percentual(totais_iaps["Recriação"], 20),
            percentual(totais_iaps["Composição"], 20),
            percentual(totais_iaps["Escuta Musical"], 20),
        ])

    if totais_demuca:
        for dominio, valor in totais_demuca.items():
            categorias.append(dominio[:18])
            valores_percentuais.append(percentual(valor, DEMUCA_DOMINIOS[dominio]["maximo"]))

    if len(categorias) < 3:
        categorias += [""] * (3 - len(categorias))
        valores_percentuais += [0] * (3 - len(valores_percentuais))

    angles = [n / float(len(categorias)) * 2 * math.pi for n in range(len(categorias))]
    valores = valores_percentuais + valores_percentuais[:1]
    angles += angles[:1]

    fig = plt.figure(figsize=(8, 8))
    fig.patch.set_facecolor("#020617")
    ax = plt.subplot(111, polar=True)
    ax.set_facecolor("#0f172a")
    ax.plot(angles, valores, linewidth=2)
    ax.fill(angles, valores, alpha=0.25)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categorias, color="white")
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(["20", "40", "60", "80", "100"], color="#cbd5e1")
    ax.set_ylim(0, 100)
    ax.set_title("Perfil Musicoterapêutico Geral (%)", color="white", pad=20, fontsize=14, fontweight="bold")
    ax.grid(color="#334155", alpha=0.8)

    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight", dpi=180, facecolor=fig.get_facecolor())
    plt.close(fig)
    buffer.seek(0)
    return buffer


# ==========================================================
# TEXTOS TÉCNICOS
# ==========================================================
def explicar_escalas(escalas_selecionadas):
    textos = []

    if "Nordoff-Robbins" in escalas_selecionadas:
        textos.append(
            "As Escalas Nordoff-Robbins são instrumentos de avaliação desenvolvidos no contexto da Musicoterapia Criativa, voltados à observação da comunicabilidade musical, da responsividade, da iniciativa, do engajamento, da reciprocidade e da organização musical do paciente no setting terapêutico. Sua aplicação permite compreender como o paciente se comunica musicalmente, como sustenta interações sonoras e como utiliza a música como recurso expressivo, relacional e regulatório."
        )

    if "IAPS" in escalas_selecionadas:
        textos.append(
            "Os IAPS, Improvisation Assessment Profiles, são instrumentos voltados à análise da improvisação clínica em musicoterapia. Eles auxiliam na observação de aspectos expressivos, interacionais, criativos, motores, cognitivos e perceptivos presentes na produção musical do paciente."
        )

    if "DEMUCA" in escalas_selecionadas:
        textos.append(
            "A DEMUCA é uma escala musicoterapêutica organizada em domínios funcionais voltados à observação de comportamentos restritivos, interação social/cognição, percepção e exploração rítmica, percepção e exploração sonora, exploração vocal e movimentação corporal com a música. Sua estrutura utiliza a classificação N = Não, P = Pouco e M = Muito. Nos comportamentos restritivos, a pontuação é invertida, pois a ausência do comportamento restritivo indica melhor funcionamento clínico. Nos demais domínios, maior pontuação indica maior presença da habilidade observada."
        )

    textos.append(
        "A escala GAS, Goal Attainment Scaling, é utilizada para estruturar metas terapêuticas individualizadas. Ela descreve níveis esperados de evolução funcional, de -2 a +2, permitindo mensurar o progresso do paciente em relação a objetivos clínicos específicos."
    )

    return "\n\n".join(textos)


def interpretar_nordoff(dados):
    if not dados:
        return ""
    total = sum(dados.values())
    nivel = classificar(total, 50)
    if nivel == "baixo":
        return "Os resultados indicam baixa disponibilidade comunicativa musical, com necessidade de maior suporte para responsividade, iniciativa, reciprocidade e sustentação da interação musical."
    if nivel == "moderado":
        return "Os resultados indicam presença de recursos comunicativos musicais em desenvolvimento, com respostas funcionais, porém ainda oscilantes em termos de iniciativa, sustentação, reciprocidade e organização da interação musical."
    return "Os resultados indicam boa comunicabilidade musical, com presença de iniciativa, responsividade, engajamento e organização relacional no fazer musical."


def interpretar_iaps(totais):
    texto = ""
    for area, valor in totals_to_order(totais).items():
        nivel = classificar(valor, 20)
        if nivel == "baixo":
            texto += f"{area}: desempenho baixo, sugerindo necessidade de intervenções estruturadas para ampliar recursos musicais, expressivos, perceptivos e relacionais.\n"
        elif nivel == "moderado":
            texto += f"{area}: desempenho moderado, indicando recursos presentes, porém ainda em processo de consolidação clínica.\n"
        else:
            texto += f"{area}: desempenho adequado, indicando boa disponibilidade funcional para esse domínio musical.\n"
    return texto


def interpretar_demuca(totais_demuca):
    texto = ""
    for dominio, valor in totais_demuca.items():
        maximo = DEMUCA_DOMINIOS[dominio]["maximo"]
        nivel = classificar(valor, maximo)
        if dominio == "Comportamentos Restritivos":
            if nivel == "baixo":
                texto += f"{dominio}: pontuação {valor}/{maximo}, classificada como baixa. Considerando a lógica invertida, sugere presença relevante de comportamentos restritivos interferindo na disponibilidade musical e relacional.\n"
            elif nivel == "moderado":
                texto += f"{dominio}: pontuação {valor}/{maximo}, classificada como moderada. Sugere presença parcial ou oscilante de comportamentos restritivos.\n"
            else:
                texto += f"{dominio}: pontuação {valor}/{maximo}, classificada como adequada. Sugere menor interferência de comportamentos restritivos no setting.\n"
        else:
            if nivel == "baixo":
                texto += f"{dominio}: pontuação {valor}/{maximo}, classificada como baixa. Indica habilidade pouco disponível ou pouco organizada no contexto sonoro-musical.\n"
            elif nivel == "moderado":
                texto += f"{dominio}: pontuação {valor}/{maximo}, classificada como moderada. Indica repertório em desenvolvimento.\n"
            else:
                texto += f"{dominio}: pontuação {valor}/{maximo}, classificada como adequada. Indica recursos funcionais importantes para o processo terapêutico.\n"
    return texto


def gerar_resumo_geral(nordoff_total, totais_iaps, totais_demuca):
    texto = (
        f"{nome if nome else 'O paciente'}, {idade} anos, com diagnóstico de {diagnostico if diagnostico else 'não informado'}, foi avaliado em musicoterapia por meio dos instrumentos selecionados. A avaliação contemplou aspectos relacionados à comunicação musical, responsividade, engajamento, interação, percepção rítmica, percepção sonora, exploração vocal, movimentação corporal, expressividade e organização musical.\n\n"
    )

    if nordoff_total is not None:
        texto += f"Na Escala Nordoff-Robbins, o desempenho geral foi classificado como {classificar(nordoff_total, 50)}. {interpretar_nordoff(nordoff_dominios)}\n\n"

    if totais_iaps:
        texto += "Nos IAPS, observou-se o seguinte perfil:\n"
        texto += interpretar_iaps(totais_iaps) + "\n"

    if totais_demuca:
        texto += "Na DEMUCA, observou-se o seguinte perfil:\n"
        texto += interpretar_demuca(totais_demuca) + "\n"

    texto += (
        "A integração dos dados sugere a necessidade de um planejamento terapêutico individualizado, fundamentado nas respostas musicais observadas e na relação entre comunicação, escuta, corpo, voz, ritmo, interação e regulação. O plano deve priorizar as habilidades em prejuízo, utilizando os recursos preservados como vias de acesso terapêutico."
    )
    return texto


def gerar_gas_automatico(prejuizos):
    metas = []
    for idx, item in enumerate(prejuizos[:3], start=1):
        habilidade = item["habilidade"]
        objetivo = item["objetivo"]
        metas.append({
            "meta": f"META {idx:02d}: {habilidade.capitalize()}",
            "-2": f"Não demonstra {habilidade} de forma funcional durante as experiências musicoterapêuticas, mesmo com mediação intensa.",
            "-1": f"Demonstra {habilidade} de forma inconsistente, com necessidade de suporte máximo, alta previsibilidade e mediação contínua do musicoterapeuta.",
            "0": objetivo,
            "+1": f"Demonstra {habilidade} com maior consistência, necessitando de suporte moderado e apresentando respostas musicais mais organizadas e intencionais.",
            "+2": f"Demonstra {habilidade} de forma funcional, espontânea e generalizável no setting musicoterapêutico, com maior autonomia e qualidade relacional."
        })
    return metas


def gerar_objetivos(prejuizos):
    return [p["objetivo"] for p in prejuizos]


def gerar_prescricao_automatica(prejuizos):
    principais = ", ".join([p["area"] for p in prejuizos[:4]])
    return (
        "Recomenda-se acompanhamento musicoterapêutico regular, com frequência mínima de 1 a 2 sessões semanais, duração média de 40 a 50 minutos, podendo ser ajustada conforme tolerância, disponibilidade atencional, perfil sensorial e resposta clínica do paciente.\n\n"
        f"O plano terapêutico deverá priorizar os domínios identificados como mais prejudicados: {principais}. As intervenções deverão ser estruturadas a partir de experiências musicais graduadas, com objetivos funcionais claros, previsibilidade, repetição terapêutica, variação progressiva de complexidade e uso de repertório significativo para o paciente.\n\n"
        "A prescrição musicoterapêutica deverá incluir improvisação clínica, recriação musical, escuta ativa, composição guiada, exploração vocal, atividades rítmicas, propostas de movimento com música e intervenções de co-regulação sonoro-musical. O setting deverá ser organizado de forma a favorecer segurança, vínculo, engajamento, iniciativa, atenção compartilhada e participação ativa."
    )


def gerar_conduta_automatizada(nordoff_total, totais_iaps, totais_demuca, prejuizos):
    texto = (
        "A partir da análise integrada dos dados avaliativos, recomenda-se a continuidade do acompanhamento musicoterapêutico com planejamento individualizado, considerando a relação entre comunicabilidade musical, expressão sonora, organização temporal, responsividade relacional, escuta, criatividade, recursos sensório-motores, exploração vocal, percepção rítmica, percepção sonora, comportamentos restritivos e possibilidades de autorregulação do paciente.\n\n"
        "As áreas em prejuízo indicam necessidade de intervenções direcionadas, musicalmente estruturadas e clinicamente graduadas. O trabalho deverá partir de experiências sonoro-musicais acessíveis, favorecendo previsibilidade, vínculo terapêutico, co-regulação, iniciativa musical, organização da resposta, sustentação atencional, comunicação não verbal, exploração vocal e integração corpo-música.\n\n"
        "Os principais focos terapêuticos identificados foram:\n"
    )

    for p in prejuizos[:5]:
        texto += f"- {p['area']} ({p['origem']}): {p['objetivo']}\n"

    texto += (
        "\nA conduta musicoterapêutica deverá integrar improvisação clínica, recriação musical, escuta ativa, composição guiada, exploração sonoro-corporal, exploração vocal, experiências rítmicas e atividades musicais significativas, sempre considerando o perfil sensorial, afetivo, cognitivo, motor, comunicativo e relacional do paciente. Recomenda-se reavaliação periódica dos objetivos terapêuticos, com ajustes conforme a evolução clínica, o engajamento nas sessões e a qualidade das respostas musicais observadas.\n\n"
        "Ao final deste parecer técnico, ressalta-se a importância de manter sessões de musicoterapia de forma regular, pois a continuidade do processo favorece a consolidação de habilidades, a ampliação da comunicação musical, a estabilidade regulatória e o desenvolvimento progressivo dos objetivos terapêuticos estabelecidos."
    )
    return texto


def gerar_estrategias(prejuizos):
    estrategias = [
        "Utilizar improvisação clínica estruturada para favorecer diálogo sonoro, turn-taking e responsividade musical.",
        "Empregar canções estruturadas para previsibilidade, organização temporal e sustentação da atenção.",
        "Realizar atividades de escuta ativa com contrastes sonoros, pausas e mediação terapêutica.",
        "Propor experiências de recriação musical com repertório significativo para o paciente.",
        "Desenvolver propostas de composição guiada para estimular autoria, simbolização e expressão emocional.",
        "Organizar o setting musicoterapêutico de forma previsível, acessível e ajustada às necessidades sensoriais e relacionais do paciente."
    ]
    for p in prejuizos:
        if "vocal" in p["habilidade"]:
            estrategias.append("Utilizar jogos vocais, vocalizações espelhadas, canções responsivas e imitação melódica para estimular expressão vocal.")
        if "rítmica" in p["habilidade"] or "temporal" in p["habilidade"]:
            estrategias.append("Aplicar padrões rítmicos simples, pulso marcado, alternância de andamento e atividades de apoio rítmico corporal/instrumental.")
        if "corporal" in p["habilidade"] or "motora" in p["habilidade"]:
            estrategias.append("Integrar movimento corporal, gestos, deslocamentos e instrumentos de fácil acesso para favorecer coordenação música-movimento.")
        if "regulação" in p["habilidade"]:
            estrategias.append("Utilizar músicas previsíveis, andamento estável, contorno melódico simples e pausas terapêuticas para co-regulação.")
    return list(dict.fromkeys(estrategias))


def referencias(escalas_escolhidas):
    refs = []
    if "Nordoff-Robbins" in escalas_escolhidas:
        refs.append("Nordoff, P., & Robbins, C. (2007). Creative Music Therapy. Gilsum: Barcelona Publishers.")
    if "IAPS" in escalas_escolhidas:
        refs.append("Bruscia, K. E. (1987). Improvisational Models of Music Therapy. Springfield: Charles C Thomas.")
    if "DEMUCA" in escalas_escolhidas:
        refs.append("DEMUCA. Escala de avaliação musicoterapêutica organizada por domínios funcionais: comportamentos restritivos, interação social/cognição, percepção rítmica, percepção sonora, exploração vocal e movimentação corporal com a música.")

    refs.extend([
        "Bruscia, K. (1998). Defining Music Therapy. Barcelona Publishers.",
        "Wigram, T., Pedersen, I. N., & Bonde, L. O. (2002). A Comprehensive Guide to Music Therapy. Jessica Kingsley Publishers.",
        "Kiresuk, T. J., & Sherman, R. E. (1968). Goal Attainment Scaling: A general method for evaluating comprehensive community mental health programs. Community Mental Health Journal."
    ])
    return refs


# ==========================================================
# WORD
# ==========================================================
def limpar_documento(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_section(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = 3
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    run.font.size = Pt(11)


def add_thumbnail_metodos(doc):
    for imagem in ["thumbnail_metodos_intensivos.png", "tumbnail_metodos_intensivos.png"]:
        if os.path.exists(imagem):
            p = doc.add_paragraph()
            p.alignment = 1
            run = p.add_run()
            run.add_picture(imagem, width=Inches(2.2))
            return


def criar_word_modelo(
    grafico_nordoff,
    grafico_iaps,
    grafico_demuca,
    grafico_radar,
    totais_iaps,
    totais_demuca,
    nordoff_total,
    resumo_geral,
    conduta_final,
    prescricao_auto,
    gas_auto,
    prejuizos,
    escalas_escolhidas
):
    try:
        doc = Document("modelo_relatorio.docx")
    except Exception:
        doc = Document()

    limpar_documento(doc)

    if metodo_intensivo in ["MIG", "TREINI"]:
        add_thumbnail_metodos(doc)

    add_title(doc, "RELATÓRIO MUSICOTERAPÊUTICO")

    add_section(doc, "IDENTIFICAÇÃO")
    tabela_id = doc.add_table(rows=0, cols=2)
    tabela_id.style = "Table Grid"

    dados_id = [
        ("Nome do paciente", nome),
        ("Data de nascimento", data_nascimento),
        ("Idade", str(idade)),
        ("Escolaridade", escolaridade),
        ("Nome dos pais ou responsáveis", responsaveis),
        ("Métodos Intensivos / Abordagem Terapêutica", metodo_intensivo),
        ("Diagnóstico", diagnostico),
        ("Data de início da intervenção", data_inicio),
        ("Data atual", data_atual),
        ("Nome do profissional", terapeuta),
        ("Registro profissional", registro),
    ]

    for campo_nome, campo_valor in dados_id:
        row = tabela_id.add_row().cells
        row[0].text = campo_nome
        row[1].text = campo_valor if campo_valor else ""

    add_section(doc, "HISTÓRIA CLÍNICA")
    add_text(doc, historia_clinica if historia_clinica else "Não informado.")

    add_section(doc, "QUEIXA PRINCIPAL / MOTIVO DO ENCAMINHAMENTO")
    add_text(doc, queixa if queixa else "Não informado.")

    add_section(doc, "OBSERVAÇÕES CLÍNICAS")
    add_text(doc, observacoes if observacoes else "Não informado.")

    add_section(doc, "AVALIAÇÕES MUSICOTERAPÊUTICAS")
    add_text(
        doc,
        f"Em {data_atual if data_atual else 'data não informada'}, foram realizadas avaliações musicoterapêuticas com base nos instrumentos selecionados, com o objetivo de coletar dados funcionais para elaboração do plano musicoterapêutico especializado."
    )

    add_section(doc, "INSTRUMENTOS AVALIATIVOS")
    add_text(doc, explicar_escalas(escalas_escolhidas))

    add_section(doc, "PREFERÊNCIAS E RECUSAS SONORO-MUSICAIS")
    add_text(doc, preferencias if preferencias else "Não informado.")

    add_section(doc, "INTERAÇÕES SONORO-MUSICAIS")
    add_text(doc, interacoes if interacoes else "Não informado.")

    add_section(doc, "RESUMO GERAL")
    add_text(doc, resumo_geral)

    if grafico_nordoff:
        add_section(doc, "Gráfico - Escala de Comunicabilidade Musical Nordoff-Robbins")
        doc.add_picture(grafico_nordoff, width=Inches(6.2))

    if totais_iaps:
        add_section(doc, "Resultados IAPS")
        tabela_iaps = doc.add_table(rows=1, cols=3)
        tabela_iaps.style = "Table Grid"
        hdr = tabela_iaps.rows[0].cells
        hdr[0].text = "Área"
        hdr[1].text = "Pontuação"
        hdr[2].text = "Classificação"

        for area, valor in totals_to_order(totais_iaps).items():
            row = tabela_iaps.add_row().cells
            row[0].text = area
            row[1].text = f"{valor}/20"
            row[2].text = classificar(valor, 20).capitalize()

        add_section(doc, "Gráfico - IAPS")
        doc.add_picture(grafico_iaps, width=Inches(6.2))

    if totais_demuca:
        add_section(doc, "Resultados DEMUCA")
        tabela_demuca = doc.add_table(rows=1, cols=3)
        tabela_demuca.style = "Table Grid"
        hdr = tabela_demuca.rows[0].cells
        hdr[0].text = "Domínio"
        hdr[1].text = "Pontuação"
        hdr[2].text = "Classificação"

        for dominio, valor in totais_demuca.items():
            maximo = DEMUCA_DOMINIOS[dominio]["maximo"]
            row = tabela_demuca.add_row().cells
            row[0].text = dominio
            row[1].text = f"{valor}/{maximo}"
            row[2].text = classificar(valor, maximo).capitalize()

        add_section(doc, "Interpretação DEMUCA")
        add_text(doc, interpretar_demuca(totais_demuca))

        add_section(doc, "Gráfico - DEMUCA")
        doc.add_picture(grafico_demuca, width=Inches(6.2))

    add_section(doc, "Gráfico - Perfil Musicoterapêutico Geral")
    doc.add_picture(grafico_radar, width=Inches(5.7))

    add_section(doc, "PLANO TERAPÊUTICO")
    add_section(doc, "Objetivos terapêuticos gerados automaticamente")
    for obj in gerar_objetivos(prejuizos):
        add_bullet(doc, obj)

    add_section(doc, "Estratégias terapêuticas sugeridas")
    for est in gerar_estrategias(prejuizos):
        add_bullet(doc, est)

    add_section(doc, "GAS / SMART GERADO AUTOMATICAMENTE")
    for meta in gas_auto:
        add_text(doc, meta["meta"])
        tabela_gas = doc.add_table(rows=1, cols=3)
        tabela_gas.style = "Table Grid"
        hdr = tabela_gas.rows[0].cells
        hdr[0].text = "Escore"
        hdr[1].text = "Descrição"
        hdr[2].text = "Objetivo / Meta prevista"

        dados = [
            ("-2", "Estado atual", meta["-2"]),
            ("-1", "Abaixo do esperado", meta["-1"]),
            ("0", "Desfecho esperado após intervenção", meta["0"]),
            ("+1", "Acima do esperado", meta["+1"]),
            ("+2", "Muito acima do esperado", meta["+2"]),
        ]

        for escore, descricao, objetivo in dados:
            row = tabela_gas.add_row().cells
            row[0].text = escore
            row[1].text = descricao
            row[2].text = objetivo

    add_section(doc, "PLANO TERAPÊUTICO E PRESCRIÇÃO GERADOS AUTOMATICAMENTE")
    add_text(doc, prescricao_auto)

    add_section(doc, "PARECER TÉCNICO / CONDUTA")
    add_text(doc, conduta_final)

    add_section(doc, "REFERÊNCIAS")
    for ref in referencias(escalas_escolhidas):
        add_bullet(doc, ref)

    doc.add_paragraph("\n\n__________________________________")
    doc.add_paragraph(terapeuta)
    doc.add_paragraph("MUSICOTERAPEUTA")
    doc.add_paragraph(registro)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# CÁLCULO GLOBAL PARA RESULTADOS E RELATÓRIO
# ==========================================================
def calcular_tudo():
    nordoff_total = sum(nordoff_dominios.values()) if nordoff_dominios else None

    totais_iaps = {}
    if iaps_improvisacao:
        totais_iaps = {
            "Improvisação": sum(iaps_improvisacao.values()),
            "Recriação": sum(iaps_recriacao.values()),
            "Composição": sum(iaps_composicao.values()),
            "Escuta Musical": sum(iaps_escuta.values()),
        }

    totais_demuca = {}
    if demuca_respostas:
        totais_demuca = calcular_demuca(demuca_respostas)

    prejuizos = identificar_prejuizos(nordoff_total, totais_iaps, totais_demuca)
    gas_auto = gerar_gas_automatico(prejuizos)
    prescricao_auto = gerar_prescricao_automatica(prejuizos)
    resumo_geral = gerar_resumo_geral(nordoff_total, totais_iaps, totais_demuca)
    conduta_final = gerar_conduta_automatizada(nordoff_total, totais_iaps, totais_demuca, prejuizos)

    grafico_nordoff = None
    if nordoff_dominios:
        grafico_nordoff = gerar_grafico_barras(
            "Escala de Comunicabilidade Musical Nordoff-Robbins",
            nordoff_dominios,
            {k: 5 for k in nordoff_dominios}
        )

    grafico_iaps = gerar_grafico_iaps(totais_iaps) if totais_iaps else None
    grafico_demuca = gerar_grafico_demuca(totais_demuca) if totais_demuca else None
    grafico_radar = gerar_grafico_radar(totais_iaps, nordoff_total, totais_demuca)

    return {
        "nordoff_total": nordoff_total,
        "totais_iaps": totais_iaps,
        "totais_demuca": totais_demuca,
        "prejuizos": prejuizos,
        "gas_auto": gas_auto,
        "prescricao_auto": prescricao_auto,
        "resumo_geral": resumo_geral,
        "conduta_final": conduta_final,
        "grafico_nordoff": grafico_nordoff,
        "grafico_iaps": grafico_iaps,
        "grafico_demuca": grafico_demuca,
        "grafico_radar": grafico_radar,
    }


# ==========================================================
# ABA 3 — RESULTADOS
# ==========================================================
with aba_resultados:
    dados = calcular_tudo()

    card_inicio("📊 Painel de análise automática", "Síntese visual dos dados coletados nas escalas selecionadas.")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Escalas selecionadas", len(escalas_escolhidas))
    with col2:
        st.metric("Áreas em foco", len(dados["prejuizos"]))
    with col3:
        st.metric("Metas GAS geradas", len(dados["gas_auto"]))
    with col4:
        st.metric("Método", metodo_intensivo if metodo_intensivo else "Não informado")
    card_fim()

    card_inicio("📈 Gráficos clínicos", "Visualização rápida do perfil musicoterapêutico.")
    if dados["grafico_nordoff"]:
        st.subheader("Nordoff-Robbins")
        st.image(dados["grafico_nordoff"], use_container_width=True)

    if dados["grafico_iaps"]:
        st.subheader("IAPS")
        st.image(dados["grafico_iaps"], use_container_width=True)

    if dados["grafico_demuca"]:
        st.subheader("DEMUCA")
        st.image(dados["grafico_demuca"], use_container_width=True)

    st.subheader("Perfil Musicoterapêutico Geral")
    st.image(dados["grafico_radar"], use_container_width=True)
    card_fim()

    card_inicio("🎯 Áreas prioritárias identificadas", "O sistema utiliza os scores para identificar prioridades clínicas e gerar objetivos automáticos.")
    for p in dados["prejuizos"]:
        st.markdown(f"**{p['area']}** — {p['origem']} | nível: `{p['nivel']}`")
        st.markdown(f"<p class='small-muted'>{p['objetivo']}</p>", unsafe_allow_html=True)
        st.markdown("---")
    card_fim()


# ==========================================================
# ABA 4 — RELATÓRIO
# ==========================================================
with aba_relatorio:
    dados = calcular_tudo()

    card_inicio("📄 Relatório clínico automatizado", "Revise os textos gerados automaticamente antes de baixar o Word.")
    st.subheader("Resumo geral")
    st.text_area("Resumo", dados["resumo_geral"], height=260)

    st.subheader("GAS / SMART automático")
    for meta in dados["gas_auto"]:
        with st.expander(meta["meta"], expanded=False):
            st.write(meta)

    st.subheader("Plano terapêutico e prescrição")
    st.text_area("Plano e prescrição", dados["prescricao_auto"], height=260)

    st.subheader("Parecer técnico / Conduta")
    st.text_area("Parecer", dados["conduta_final"], height=300)
    card_fim()

    card_inicio("⬇️ Gerar documento Word", "Clique para gerar e baixar o relatório final em formato .docx.")

    if st.button("🚀 Preparar relatório clínico"):
        word = criar_word_modelo(
            grafico_nordoff=dados["grafico_nordoff"],
            grafico_iaps=dados["grafico_iaps"],
            grafico_demuca=dados["grafico_demuca"],
            grafico_radar=dados["grafico_radar"],
            totais_iaps=dados["totais_iaps"],
            totais_demuca=dados["totais_demuca"],
            nordoff_total=dados["nordoff_total"],
            resumo_geral=dados["resumo_geral"],
            conduta_final=dados["conduta_final"],
            prescricao_auto=dados["prescricao_auto"],
            gas_auto=dados["gas_auto"],
            prejuizos=dados["prejuizos"],
            escalas_escolhidas=escalas_escolhidas
        )

        st.session_state.word = word
        st.session_state.relatorio_gerado = True
        st.success("Relatório preparado com sucesso.")

    if st.session_state.get("relatorio_gerado", False):
        nome_arquivo = f"relatorio_{nome.replace(' ', '_') if nome else 'paciente'}.docx"
        st.download_button(
            "📄 Baixar relatório em Word",
            data=st.session_state.word,
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    card_fim()
